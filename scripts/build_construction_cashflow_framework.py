from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("outputs/construction-cashflow-framework/RESLU-Construction-Job-Cycle-and-Cash-Flow-Framework.docx")

INK = "242320"
CHARCOAL = "4A4945"
SAND = "9A7B55"
CREAM = "F5F1E8"
PALE = "ECE7DD"
LIGHT = "F7F5F0"
MID = "D9D2C5"
WHITE = "FFFFFF"
GREEN = "E7EFE5"
AMBER = "F5EBD7"


def set_font(run, size=11, bold=False, italic=False, color=INK, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=MID, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def apply_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def format_cell_text(cell, *, bold=False, color=INK, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    for p in cell.paragraphs:
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        for r in p.runs:
            set_font(r, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths_dxa, *, header_fill=PALE, font_size=9.2, aligns=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        header.cells[idx].text = text
        shade(header.cells[idx], header_fill)
        format_cell_text(header.cells[idx], bold=True, size=font_size, align=(aligns or {}).get(idx, WD_ALIGN_PARAGRAPH.LEFT))
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            if row_index % 2 == 1:
                shade(cells[idx], LIGHT)
            format_cell_text(cells[idx], size=font_size, align=(aligns or {}).get(idx, WD_ALIGN_PARAGRAPH.LEFT))
    apply_table_geometry(table, widths_dxa)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text, *, bold_lead=None, italic=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    if bold_lead and text.startswith(bold_lead):
        a = p.add_run(bold_lead)
        set_font(a, bold=True)
        b = p.add_run(text[len(bold_lead):])
        set_font(b, italic=italic)
    else:
        r = p.add_run(text)
        set_font(r, italic=italic)
    return p


def add_callout(doc, label, text, fill=CREAM):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True, color=SAND)
    r = p.add_run(text)
    set_font(r)
    set_table_borders(table, color=MID, size=5)
    apply_table_geometry(table, [9360])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_checkbox_table(doc, items):
    rows = [["☐", item, ""] for item in items]
    return add_table(
        doc,
        ["", "Decision to confirm", "Notes / agreed rule"],
        rows,
        [430, 4850, 4080],
        font_size=9.5,
        aligns={0: WD_ALIGN_PARAGRAPH.CENTER},
    )


def set_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for level, size, before, after in ((1, 16, 15, 7), (2, 13, 11, 5), (3, 11.5, 8, 4)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(SAND if level < 3 else CHARCOAL)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("RESLU  |  CONSTRUCTION CASH-FLOW FRAMEWORK")
    set_font(r, size=8.2, bold=True, color=CHARCOAL)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("Working draft  •  ")
    set_font(r, size=8, color=CHARCOAL)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)


def build():
    doc = Document()
    doc.settings.odd_and_even_pages_header = False
    set_styles(doc)
    props = doc.core_properties
    props.title = "RESLU Construction Job Cycle and Cash-Flow Framework"
    props.subject = "Editable workshop framework for construction stages, trade packages and payment timing"
    props.author = "RESLU"
    props.keywords = "construction, cash flow, job cycle, trades, payments, FF&E"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("WORKING FRAMEWORK")
    set_font(r, size=9, bold=True, color=SAND)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    r = title.add_run("Construction Job Cycle & Cash-Flow Framework")
    set_font(r, size=25, bold=True, color=INK, name="Aptos Display")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(15)
    r = subtitle.add_run("An editable working document for agreeing RESLU’s construction stages, trade packages and standard payment rules before changing Spec.")
    set_font(r, size=11.5, color=CHARCOAL)

    meta = add_table(
        doc,
        ["Purpose", "Status", "Prepared"],
        [["Workshop and validation", "Draft — not system logic", date(2026, 8, 21).strftime("%d %B %Y")]],
        [3900, 2700, 2760],
        header_fill=CREAM,
        font_size=9.5,
    )

    add_callout(
        doc,
        "Recommended method",
        "Agree the stages first, map trade packages second, and define payment rules third. Then test the complete framework against one real construction job. Only rules that survive that test should become automated fields in Spec.",
        fill=GREEN,
    )

    add_heading(doc, "1. What this framework must solve", 1)
    add_body(doc, "The current forecast requires too many manual dates because the cost, the work stage and the payment event are not consistently connected. This framework separates those concepts so a programme movement can update forecast timing without changing the underlying approved budget.")
    add_table(
        doc,
        ["Layer", "Question it answers", "Proposed source of truth"],
        [
            ["Job cycle", "When does the work occur?", "The project construction programme / Timeline"],
            ["Trade package", "Who performs or supplies the work?", "A controlled RESLU trade-package list"],
            ["Cost record", "What are we paying for and how much?", "Estimate line, trade package or FF&E item"],
            ["Payment structure", "What event makes cash leave RESLU?", "A standard payment template with permitted overrides"],
            ["Actual", "What was invoiced and paid?", "Supplier invoice, payment record and Xero"],
        ],
        [1700, 3300, 4360],
    )

    add_heading(doc, "2. Proposed construction job cycle", 1)
    add_body(doc, "This is a starting sequence for review. The stage names should be broad enough to work across projects, while trade-specific detail sits in the trade-package mapping that follows.")
    cycle_rows = [
        ["01", "Pre-construction", "Contract, approvals, programme, procurement planning", "Programme approved"],
        ["02", "Site establishment", "Access, protection, temporary services, safety setup", "Site ready"],
        ["03", "Demolition & strip-out", "Demolition, disposal, make-safe and exposure", "Demolition complete"],
        ["04", "Structural & framing", "Footings, steel, wall/roof framing and structural sign-off", "Frame approved"],
        ["05", "External envelope", "Roofing, windows, doors, brickwork, cladding and sealing", "Build sealed"],
        ["06", "Services rough-in / first fix", "Concealed plumbing, electrical, HVAC and data services", "Rough-in complete"],
        ["07", "Internal linings & waterproofing", "Insulation, plasterboard, screed beds and membranes", "Linings/waterproofing complete"],
        ["08", "Internal finishes", "Tiling, flooring and applied feature finishes", "Finishes complete"],
        ["09", "Joinery & fixed elements", "Kitchen, vanities, wardrobes and other fixed joinery", "Joinery installed"],
        ["10", "Fit-off / second fix", "Visible fixtures, fittings, appliances, doors, trims and hardware", "Fit-off complete"],
        ["11", "Painting & final detail", "Painting, touch-ups, sealants, detailing and final clean", "Ready for completion inspection"],
        ["12", "Practical completion", "Inspection, defects, certificates and rectification", "Practical completion"],
        ["13", "Handover & close-out", "Client walkthrough, manuals, warranties and final account", "Job closed"],
    ]
    add_table(doc, ["ID", "Job-cycle stage", "Typical scope", "Stage completion event"], cycle_rows, [520, 2050, 4300, 2490], font_size=8.7, aligns={0: WD_ALIGN_PARAGRAPH.CENTER})

    add_heading(doc, "3. Trade packages within the cycle", 1)
    add_body(doc, "Avoid generic labels such as “first fix” on their own. The cash-flow and reporting logic should identify both the trade and the stage, for example Carpentry — first fix or Plumbing — fit-off.")
    trade_rows = [
        ["Demolition", "Demolition & strip-out", "Strip-out, removal, disposal, make-safe", "Demolition complete"],
        ["Carpentry — first fix", "Structural & framing", "Framing, structural timber, concealed backing and support", "Measured progress or stage completion"],
        ["Carpentry — second fix", "Fit-off / second fix", "Doors, jambs, skirtings, architraves, trims and hardware", "Measured progress or installation complete"],
        ["Plumbing — rough-in", "Services rough-in / first fix", "Concealed water, waste and gas services", "Rough-in inspection / completion"],
        ["Plumbing — fit-off", "Fit-off / second fix", "Tapware, sanitaryware, fixtures and commissioning", "Fit-off completion"],
        ["Electrical — rough-in", "Services rough-in / first fix", "Cabling, boxes, switchboard and concealed services", "Rough-in inspection / completion"],
        ["Electrical — fit-off", "Fit-off / second fix", "Lights, switches, outlets, appliances and testing", "Fit-off completion"],
        ["HVAC — rough-in", "Services rough-in / first fix", "Ducts, pipework, cabling and concealed plant connections", "Rough-in complete"],
        ["HVAC — fit-off", "Fit-off / second fix", "Grilles, controls, commissioning and balancing", "Commissioned"],
        ["Waterproofing", "Internal linings & waterproofing", "Screed interface, membranes and certification", "Certificate / stage complete"],
        ["Tiling", "Internal finishes", "Wall and floor tiling, grout and seal", "Area complete"],
        ["Flooring", "Internal finishes", "Timber, resilient, carpet and associated preparation", "Area complete"],
        ["Joinery", "Joinery & fixed elements", "Manufacture, delivery and installation of fixed joinery", "Installed / defects recorded"],
        ["Painting", "Painting & final detail", "Preparation, coats, touch-ups and final detail", "Painting complete"],
        ["Appliances & FF&E", "Fit-off / second fix", "Ordered products, delivery and installation", "Delivered or installed, depending on payment rule"],
        ["External works", "External works / project-specific", "Landscaping, paving, fencing, gates and drainage", "Package complete"],
    ]
    add_table(doc, ["Trade package", "Default job stage", "Included scope", "Default progress event"], trade_rows, [2110, 2190, 3260, 1800], font_size=8.35)

    add_heading(doc, "4. Standard payment structures", 1)
    add_body(doc, "A work stage describes when the work happens. A payment structure describes when cash leaves RESLU. They must remain separate: a basin may be installed at Plumbing — fit-off but paid for when ordered weeks earlier.")
    payment_rows = [
        ["P1", "Trade stage claim", "100%", "Stage or measured work completed", "Invoice date + supplier terms", "Labour and straightforward trade packages"],
        ["P2", "Deposit + completion", "50% / 50%", "Acceptance or mobilisation / package completion", "Trigger date + supplier terms", "Small fabricated packages and some subcontractors"],
        ["P3", "Deposit + pre-delivery balance", "50% / 50%", "Order placed / before dispatch", "Order date / ETA less agreed offset", "Custom products and supplier orders"],
        ["P4", "Progress claims", "Custom, totals 100%", "Measured progress milestones", "Each milestone + supplier terms", "Carpentry, joinery and larger trade packages"],
        ["P5", "Pay on order", "100%", "Purchase order accepted", "Order date + supplier terms", "Stock items or suppliers requiring full payment"],
        ["P6", "Pay on delivery", "100%", "Goods delivered", "Delivery date + supplier terms", "Products supplied on account"],
        ["P7", "Pay on installation", "100%", "Installation completed", "Installation date + supplier terms", "Supply-and-install packages"],
        ["P8", "Custom exception", "Custom, totals 100%", "Explicit approved milestones", "Explicit rule per milestone", "Only where no standard template fits"],
    ]
    add_table(doc, ["Code", "Template", "Split", "Forecast trigger", "Expected cash date", "Typical use"], payment_rows, [560, 1740, 1000, 2220, 2100, 1740], font_size=8.15, aligns={0: WD_ALIGN_PARAGRAPH.CENTER, 2: WD_ALIGN_PARAGRAPH.CENTER})

    add_callout(doc, "Control rule", "Every payment template must total 100%. A matched supplier invoice replaces the forecast amount and due date; a paid transaction replaces the invoice in actual cash flow.", fill=AMBER)

    add_heading(doc, "5. Proposed forecasting logic", 1)
    logic_rows = [
        ["1", "Start with the approved cost", "Estimate, purchase order or current approved FF&E cost"],
        ["2", "Assign the trade package", "For example Plumbing — fit-off"],
        ["3", "Link to the project programme", "Use the actual Timeline stage or milestone, not a copied manual date"],
        ["4", "Apply the payment template", "Split the approved cost across payment milestones"],
        ["5", "Calculate forecast cash dates", "Trigger date, lead time, offset and supplier payment terms"],
        ["6", "Replace forecast with better evidence", "Purchase order → invoice → paid transaction"],
        ["7", "Retain baseline for comparison", "Programme and forecast can move without rewriting the approved baseline"],
    ]
    add_table(doc, ["Step", "Action", "Rule"], logic_rows, [620, 2730, 6010], font_size=9.2, aligns={0: WD_ALIGN_PARAGRAPH.CENTER})

    add_heading(doc, "6. Worked construction examples", 1)
    add_body(doc, "Use these rows as a discussion starter. Replace them with one recent RESLU job and compare the calculated timing with what actually occurred.")
    example_rows = [
        ["Framing labour", "$28,000", "Carpentry — first fix", "Structural & framing", "P4", "30% mobilisation; 50% frame progress; 20% frame approved"],
        ["Internal doors & trim", "$14,500", "Carpentry — second fix", "Fit-off / second fix", "P3", "50% on order; 50% before delivery"],
        ["Plumbing rough-in labour", "$11,000", "Plumbing — rough-in", "Services rough-in", "P1", "100% when rough-in is complete, then supplier terms"],
        ["Tapware package", "$9,800", "Plumbing — fit-off", "Fit-off / second fix", "P5", "100% paid when ordered; work stage remains fit-off"],
        ["Custom joinery", "$52,000", "Joinery", "Joinery & fixed elements", "P4", "10% design; 40% manufacture; 40% before delivery; 10% completion"],
        ["Electrical fit-off", "$8,500", "Electrical — fit-off", "Fit-off / second fix", "P1", "100% on completion, then supplier terms"],
    ]
    add_table(doc, ["Cost", "Example value", "Trade package", "Work stage", "Template", "Illustrative milestones"], example_rows, [1510, 1050, 1700, 1770, 780, 2550], font_size=8.35, aligns={1: WD_ALIGN_PARAGRAPH.RIGHT, 4: WD_ALIGN_PARAGRAPH.CENTER})

    add_heading(doc, "7. Decisions to make before Spec changes", 1)
    add_checkbox_table(doc, [
        "Confirm the final construction stage names and order.",
        "Confirm whether External works is a standard stage or only added when relevant.",
        "Confirm the complete list of trade packages and who owns each mapping.",
        "Agree whether Carpentry first fix and Carpentry second fix are separate packages on every job.",
        "Confirm the default payment template for each trade package.",
        "Agree the default supplier payment terms and whether terms can vary by supplier.",
        "Agree which event drives each milestone: order, stage start, stage completion, delivery or installation.",
        "Agree how lead time changes the order and deposit forecast dates.",
        "Confirm how variations enter the baseline and current forecast.",
        "Confirm when supplier invoices replace forecast milestones.",
        "Confirm the minimum information required before a cost is allowed into the forecast.",
        "Test the framework against one completed job and one current construction job.",
    ])

    add_heading(doc, "8. Minimum data needed in Spec", 1)
    add_body(doc, "This section is deliberately a future implementation brief, not an instruction to change the system yet.")
    fields_rows = [
        ["Cost source", "Estimate line, trade package or FF&E item", "Already mostly available"],
        ["Trade package", "Controlled RESLU list", "New shared classification"],
        ["Programme link", "Construction stage or milestone ID", "Reuse Timeline; do not duplicate dates"],
        ["Payment template", "P1–P8 or later approved set", "New standard selector"],
        ["Payment milestones", "Percentage, trigger, offset and terms", "Generated from template; override by exception"],
        ["Evidence state", "Forecast, committed, invoiced or paid", "Allows better evidence to supersede forecast"],
        ["Baseline reference", "Approved estimate / contract version", "Retains variance history"],
    ]
    add_table(doc, ["Field", "Meaning", "Design intent"], fields_rows, [1900, 3500, 3960], font_size=9.1)

    add_heading(doc, "9. Workshop notes", 1)
    for prompt in (
        "Stage changes:",
        "Trade-package changes:",
        "Payment-template changes:",
        "Exceptions we see repeatedly:",
        "Rules approved for a real-job test:",
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(prompt)
        set_font(r, bold=True, color=CHARCOAL)
        for _ in range(2):
            line = doc.add_paragraph("________________________________________________________________________________")
            line.paragraph_format.space_after = Pt(2)
            for run in line.runs:
                set_font(run, size=9, color=MID)

    add_callout(doc, "Exit criterion", "Do not automate the framework until a real-job test produces a cash forecast that the team would have trusted without manual re-dating.", fill=GREEN)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
